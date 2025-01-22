import json
import re
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import os
from datetime import datetime
from docx import Document
import sqlparse

def parse_sql_file(sql_file_path):

    with open(sql_file_path, 'r', encoding='utf-8') as sql_file:
        sql_content = sql_file.read()

    # sqlparse解析SQL
    parsed = sqlparse.parse(sql_content)

    tables= []

    for statement in parsed:
        # 将statement转换为一个单行字符串，去掉首尾的空格和换行符
        statement_str = str(statement).replace('\n', ' ').strip()

        # 如果是 CREATE TABLE 语句
        if statement_str.startswith('CREATE TABLE'):
            table_info = parse_create_table(statement_str)
            if table_info:
                tables.append(table_info)

    return tables

def parse_create_table(statement_str):
    # 匹配 CREATE TABLE 语句中的表名
    create_table_pattern = re.compile(r'CREATE TABLE\s+`?(\w+)`?\s*\(')
    match = create_table_pattern.search(statement_str)

    if match:
        # 提取 table_name
        table_name_match = re.search(r'CREATE TABLE\s+`?(\w+)`?', statement_str.strip(), re.IGNORECASE)
        table_name = table_name_match.group(1) if table_name_match else None

        # 清理输入，去掉 CREATE TABLE 部分
        statement_str = re.sub(r'CREATE TABLE.*?\(', '', statement_str, flags=re.S).strip()

        # 去除表格结构结尾部分（PRIMARY KEY, ENGINE 等）
        statement_str = re.sub(r'\)\s*ENGINE.*', ')', statement_str, flags=re.S).strip()

        # 处理字段定义部分
        column_pattern = re.compile(
            r'`?(\w+)`?\s+(\w+)(\((\d+)\))?(\s+NOT NULL)?'
)

        columns = []
        for part in statement_str.split(','):
            column_match = column_pattern.match(part.strip())
            if column_match:
                column_name = column_match.group(1).strip('`')  # 移除反引号
                column_type = column_match.group(2).strip()
                column_length = column_match.group(4)
                not_null = 'NOT NULL' in part

                # 手动查询 DEFAULT 和 COMMENT
                column_def = match.group(0)
                default_value = None
                comment_value = None

                # 查询 DEFAULT
                default_match = re.search(r'\s+DEFAULT\s+([^\s,]+)', part.strip())
                if default_match:
                    default_value = default_match.group(1)

                # 查询 COMMENT
                comment_match = re.search(r'\s+COMMENT\s+\'([^\']+)\'', part.strip())
                if comment_match:
                    comment_value = comment_match.group(1)

                columns.append({
                    'column_name': column_name,
                    'column_type': column_type,
                    'column_length': column_length,
                    'not_null': not_null,
                    'default_value': default_value,
                    'comment': comment_value
                })

        return {
            'table_name': table_name,
            'columns': columns
        }

    return None


def create_word_table(tables):
    # 创建一个新的 Word 文档
    doc = Document()

    # 遍历每个表
    for table in tables:
        # 添加表名
        doc.add_heading(table['table_name'], level=1)

        # 添加表格
        if table['columns']:
            # 创建表格：列数 = 5 (name, type, length, nullability, comment)
            word_table = doc.add_table(rows=1, cols=5)
            word_table.style = 'Table Grid'  # 设置表格样式

            # 填充表头
            headers = word_table.rows[0].cells
            headers[0].text = '字段名称'
            headers[1].text = '类型'
            headers[2].text = '长度'
            headers[3].text = '是否必填'
            headers[4].text = '备注'

            # 填充数据
            for column in table['columns']:
                column_name = column.get('column_name', '')
                column_type = column.get('column_type', '')

                # 使用正则表达式检查是否为主键或索引
                isPrimaryKey = (re.match(r'(PRIMARY|primary)', column_name) and
                                re.match(r'(KEY|key)', column_type))

                # 如果是主键或索引，跳过当前列
                if isPrimaryKey:
                    continue

                row_cells = word_table.add_row().cells
                row_cells[0].text = column.get('column_name', '')
                row_cells[1].text = column.get('column_type', '')
                row_cells[2].text = str(column.get('column_length', ''))
                row_cells[3].text = '是' if column.get('not_null', False) else '不是'
                row_cells[4].text = column.get('comment') if column.get('comment') is not None else ''

    # 保存文档
    word_file_path = os.path.join(os.path.expanduser('~'), 'Desktop',
                                  f"converted_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx")
    doc.save(word_file_path)

    # 显示转换结果
    messagebox.showinfo("完成", f"转换完成！文件已保存至: {word_file_path}")

# 按钮回调函数
def on_select_file():
    file_path = filedialog.askopenfilename(filetypes=[("SQL Files", "*.sql")])
    if file_path:
        entry_file_path.delete(0, tk.END)
        entry_file_path.insert(0, file_path)


def on_convert():
    sql_file_path = entry_file_path.get()
    if not sql_file_path:
        messagebox.showerror("错误", "请先选择一个SQL文件")
        return

    if not os.path.exists(sql_file_path):
        messagebox.showerror("错误", "选择的文件不存在")
        return

    # 调用转换函数
    tables = parse_sql_file(sql_file_path)
    create_word_table(tables)
    json_result = json.dumps(tables, indent=4, ensure_ascii=False)
    print(json_result)



if __name__ == '__main__':

    # 创建主窗口
    root = tk.Tk()
    root.title("SQL文件转换为Word")

    # UI
    label = tk.Label(root, text="选择一个SQL文件:")
    label.pack(padx=10, pady=5)

    entry_file_path = tk.Entry(root, width=50)
    entry_file_path.pack(padx=10, pady=5)

    button_select_file = tk.Button(root, text="选择文件", command=on_select_file)
    button_select_file.pack(padx=10, pady=5)

    button_convert = tk.Button(root, text="转换", command=on_convert)
    button_convert.pack(padx=10, pady=20)

    # 运行应用
    root.mainloop()
